import sqlite3
from docx import Document

doc_path = "User_info.docx"

# Function to clear the document
def clear_document(doc_path):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    for table in doc.tables:
        t = table._element
        t.getparent().remove(t)
    doc.save(doc_path)

# Function to add a title
def add_title(doc_path, title_text):
    doc = Document(doc_path)
    title = doc.add_paragraph(title_text)
    title.style = 'Title'
    doc.save(doc_path)

# Function to fetch users from the database
def fetch_users_from_db(db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("SELECT id, first_name, email, daily_protein_goal, daily_calorie_goal FROM User")
    rows = cursor.fetchall()
    conn.close()
    return rows

# Function to fetch meals from the database
def fetch_meals_from_db(db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("SELECT user_id, protein, calories, date FROM Meal")
    rows = cursor.fetchall()
    conn.close()
    return rows

# Function to add user login information
def add_user_logins(doc_path):
    doc = Document(doc_path)
    doc.add_heading(f"User Login Data", level=1)
    table_header = ["ID", "Name", "Email", "Protein goal", "Calorie goal"]
    table = doc.add_table(rows=1, cols=len(table_header))
    for i in range(len(table_header)):
        table.rows[0].cells[i].text = table_header[i]
    doc.save(doc_path)

# Function to add a user meal section
def add_user_meals(doc_path, user_name, user_id):
    doc = Document(doc_path)
    for text in doc.paragraphs:
        if f"ID: {user_id}" in text.text or user_name in text.text:
            raise IndexError(f"User {user_name} already exists, ID: {user_id}")
    doc.add_heading(f"Meals for: {user_name} , ID: {user_id}", level=1)
    table_header = ["Protein", "Calories", "Date"]
    table = doc.add_table(rows=1, cols=len(table_header))
    for i in range(len(table_header)):
        table.rows[0].cells[i].text = table_header[i]
    doc.save(doc_path)

# Function to add a row to a table
def add_row_to_table(doc_path, table_number, row_content):
    doc = Document(doc_path)
    if table_number >= len(doc.tables):
        raise IndexError(f"Table number {table_number} is out of range. This document contains {len(doc.tables)} tables.")
    table = doc.tables[table_number]
    new_row_cells = table.add_row().cells
    if len(row_content) != len(new_row_cells):
        raise ValueError(f"Row content length ({len(row_content)}) does not match the number of columns in the table ({len(new_row_cells)}).")
    for row in table.rows:
        for cell in row.cells:
            if row_content[1] in cell.text and row_content[2] in cell.text:
                raise ValueError(f"There is already a user with this information {row_content}, table num: {table_number}, row cnt:{cell.text}")
    for i in range(len(new_row_cells)):
        new_row_cells[i].text = row_content[i]
    doc.save(doc_path)

# Clear the document
clear_document(doc_path)

# Add a title to the cleared document
add_title(doc_path, "User Information Database")

# Fetch users and meals from the database
db_path = 'Flask website (data base)\instance\database.db'
users = fetch_users_from_db(db_path)
meals = fetch_meals_from_db(db_path)

user_login_list = [[str(user[0]), user[1], user[2], str(user[3]), str(user[4])] for user in users]
user_meals_list = [[meal[0], str(meal[1]), str(meal[2]), meal[3]] for meal in meals]

# Check start-up procedures and add user logins if necessary
def check_start_up_procedures(doc_path):
    document = Document(doc_path)
    if len(document.tables) < 1:
        add_user_logins(doc_path)
        for user in user_login_list:
                add_row_to_table(doc_path, 0, user)
            

check_start_up_procedures(doc_path)

# Add user meal sections and meal data
for user in user_login_list:
    try:
        add_user_meals(doc_path, user_name=user[1], user_id=user[0])
    except IndexError:
        pass

for meal in user_meals_list:
    try:
        add_row_to_table(doc_path, table_number=meal[0], row_content=meal[1:])
    except IndexError:
        pass

  